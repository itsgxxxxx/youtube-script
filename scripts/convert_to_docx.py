#!/usr/bin/env python3
"""Convert a Markdown transcript into a .docx Word document."""

import argparse
import os
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt
except ImportError:
    print("Error: python-docx not installed. Run: pip install python-docx")
    sys.exit(1)


def clean_text(text: str) -> str:
    """Remove basic Markdown syntax while preserving readable text."""
    text = re.sub(r"!\[.*?\]\(.*?\)", "", text)
    text = re.sub(r"\[(.*?)\]\(.*?\)", r"\1", text)
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"\*(.*?)\*", r"\1", text)
    text = re.sub(r"`(.*?)`", r"\1", text)
    text = re.sub(r"^#+\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^---+$", "", text, flags=re.MULTILINE)
    return text.strip()


def parse_markdown(md_content: str) -> list[tuple]:
    """Parse a transcript-flavored Markdown document into renderable blocks."""
    sections = []

    for raw_line in md_content.splitlines():
        line = raw_line.strip()
        if not line:
            sections.append(("blank",))
            continue

        if line.startswith("# ") and not line.startswith("## "):
            sections.append(("title", clean_text(line)))
            continue

        if line.startswith("## "):
            sections.append(("heading2", clean_text(line)))
            continue

        if line.startswith("### "):
            sections.append(("heading3", clean_text(line)))
            continue

        speaker_match = re.match(r"\*\*\[(\d{1,2}:\d{2}(?::\d{2})?)\]\s*(.*?)\*\*:\s*(.*)", line)
        if speaker_match:
            timestamp, speaker, content = speaker_match.groups()
            sections.append(("dialogue", timestamp, speaker, clean_text(content)))
            continue

        tutorial_match = re.match(r"\*\*\[(\d{1,2}:\d{2}(?::\d{2})?)\]\*\*\s*(.*)", line)
        if tutorial_match:
            timestamp, content = tutorial_match.groups()
            sections.append(("timestamped_text", timestamp, clean_text(content)))
            continue

        if line.startswith("- "):
            sections.append(("bullet", clean_text(line[2:])))
            continue

        if line.startswith("**") and "**:" in line:
            sections.append(("meta", clean_text(line)))
            continue

        sections.append(("text", clean_text(line)))

    return sections


def add_run(paragraph, text: str, bold: bool = False, font_size: int | None = None) -> None:
    run = paragraph.add_run(text)
    run.bold = bold
    if font_size:
        run.font.size = Pt(font_size)


def generate_docx(md_file: str, output_file: str | None = None) -> str:
    with open(md_file, "r", encoding="utf-8") as handle:
        md_content = handle.read()

    sections = parse_markdown(md_content)
    document = Document()

    for section in document.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    title_written = False

    for section in sections:
        section_type = section[0]

        if section_type == "blank":
            document.add_paragraph("")
            continue

        if section_type == "title":
            paragraph = document.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(paragraph, section[1], bold=True, font_size=20)
            title_written = True
            continue

        if section_type == "heading2":
            document.add_heading(section[1], level=1)
            continue

        if section_type == "heading3":
            document.add_heading(section[1], level=2)
            continue

        if section_type == "meta":
            paragraph = document.add_paragraph()
            add_run(paragraph, section[1], font_size=10)
            continue

        if section_type == "bullet":
            document.add_paragraph(section[1], style="List Bullet")
            continue

        if section_type == "dialogue":
            timestamp, speaker, content = section[1], section[2], section[3]
            paragraph = document.add_paragraph()
            add_run(paragraph, f"[{timestamp}] {speaker}: ", bold=True)
            add_run(paragraph, content)
            continue

        if section_type == "timestamped_text":
            timestamp, content = section[1], section[2]
            paragraph = document.add_paragraph()
            add_run(paragraph, f"[{timestamp}] ", bold=True)
            add_run(paragraph, content)
            continue

        document.add_paragraph(section[1])

    if not title_written:
        title = Path(md_file).stem
        paragraph = document.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(paragraph, title, bold=True, font_size=20)

    if not output_file:
        output_file = os.path.splitext(md_file)[0] + ".docx"

    document.save(output_file)
    print(f"Word document generated: {output_file}")
    return output_file


def main() -> None:
    parser = argparse.ArgumentParser(description="Convert Markdown transcript to Word (.docx)")
    parser.add_argument("md_file", help="Input Markdown file")
    parser.add_argument("-o", "--output", help="Output .docx file")
    args = parser.parse_args()

    if not os.path.exists(args.md_file):
        print(f"Error: File not found: {args.md_file}")
        sys.exit(1)

    generate_docx(args.md_file, args.output)


if __name__ == "__main__":
    main()
